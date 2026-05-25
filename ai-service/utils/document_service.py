from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import logging
import re

logger = logging.getLogger(__name__)


def apply_run_style(run, font_size=12, bold=None):
    """
    Surgically applies typography styling to a run while preserving color and other styles.
    """
    run.font.name = 'Bookman Old Style'
    run.font.size = Pt(font_size)
    if bold is not None:
        run.bold = bold
    
    # Global font mapping enforcement via XML to ensure consistency across all renderers
    r = run._element.get_or_add_rPr()
    rFonts = r.get_or_add_rFonts()
    for attr in ['w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs']:
        rFonts.set(qn(attr), 'Bookman Old Style')


def safe_update_runs(paragraph, new_text):
    """
    Updates paragraph text by modifying existing runs to preserve formatting.
    """
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    # Update first run and clear others to maintain style of the first run
    paragraph.runs[0].text = new_text
    for i in range(1, len(paragraph.runs)):
        paragraph.runs[i].text = ""


def apply_academic_formatting(paragraph, is_heading=False):
    """
    Applies standard academic styling to non-table paragraphs without destroying existing styles.
    """
    # Global Spacing Rule: 1.5 for ALL paragraphs
    paragraph.paragraph_format.line_spacing = 1.5
    
    # Alignment Protection Logic:
    # If alignment is missing (None) or broken (Right), apply specific defaults.
    # Preserve if already defined (e.g., Center, Justify, Left).
    if paragraph.alignment in [None, WD_ALIGN_PARAGRAPH.RIGHT]:
        if is_heading:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    font_size = 14 if is_heading else 12
    is_bold = True if is_heading else False
    
    if is_heading:
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)

    # Apply style to all existing runs
    for run in paragraph.runs:
        apply_run_style(run, font_size=font_size, bold=is_bold)


def process_label_value_paragraph(paragraph):
    """
    Surgically formats Label : Value patterns without paragraph.clear().
    """
    text = paragraph.text
    if ":" not in text:
        for run in paragraph.runs:
            apply_run_style(run, font_size=12, bold=False)
        return

    parts = text.split(":", 1)
    label_part = parts[0].strip() + " :"
    value_part = " " + parts[1].strip()

    if not paragraph.runs:
        paragraph.add_run("")

    # Update first run for label
    first_run = paragraph.runs[0]
    original_color = first_run.font.color.rgb
    first_run.text = label_part
    apply_run_style(first_run, font_size=12, bold=True)
    if original_color:
        first_run.font.color.rgb = original_color

    # Update or add second run for value
    if len(paragraph.runs) > 1:
        second_run = paragraph.runs[1]
        second_run.text = value_part
        apply_run_style(second_run, font_size=12, bold=False)
        # Clear any extra runs
        for i in range(2, len(paragraph.runs)):
            paragraph.runs[i].text = ""
    else:
        v_run = paragraph.add_run(value_part)
        apply_run_style(v_run, font_size=12, bold=False)


def fill_reflective_journal(template_path: str, output_path: str, data: dict) -> str:
    doc = Document(template_path)

    # Standardize Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Data mapping according to Input Contract
    # Hard Locked Fields (from user input)
    student_details = data.get("student_details", {})
    gen_content = data.get("generated_content", {})
    
    mappings = {
        "name of the student": student_details.get("student_name", ""),
        "academic year": student_details.get("academic_year", ""),
        "student registration number": student_details.get("registration_number", ""),
        "registration number": student_details.get("registration_number", ""),
        "year & term": student_details.get("year_term", ""),
        "year and term": student_details.get("year_term", ""),
        "study level": student_details.get("study_level", ""),
        "class & section": student_details.get("class_section", ""),
        "class and section": student_details.get("class_section", ""),
        "course": student_details.get("course_name", ""),
        "instructor": student_details.get("instructor", ""),
        "assessment": student_details.get("assessment", ""),
        "date": data.get("date", ""),
        "journal entry topic": data.get("journal_topic", ""),
        "experience": gen_content.get("experience", ""),
        "feelings": gen_content.get("feelings", ""),
        "learning": gen_content.get("learning", ""),
        "application": gen_content.get("application", ""),
        "conclusion": gen_content.get("conclusion", ""),
    }

    # 1. Fill data (Maintain mapping logic with strict leakage protection)
    priority_keys = [
        "journal entry topic", "name of the student", "academic year",
        "student registration number", "registration number", "year & term",
        "year and term", "study level", "class & section", "class and section",
        "course", "instructor", "assessment", "date", "experience", "feelings", "learning",
        "application", "conclusion"
    ]

    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell_text_lower = cell.text.lower().strip()
                
                for key in priority_keys:
                    val = mappings[key]
                    # Strict mapping rule: 
                    # 1. If it's the topic field, only the topic key can match.
                    # 2. Otherwise, the key must be in the cell text.
                    is_match = False
                    if "journal entry topic" in cell_text_lower:
                        if key == "journal entry topic":
                            is_match = True
                        else:
                            continue # Prevent "learning" from matching "Machine Learning" in topic
                    elif key in cell_text_lower:
                        is_match = True
                        
                    if is_match:
                        if ":" in cell.text:
                            # Safely update label:value in same cell
                            for para in cell.paragraphs:
                                if ":" in para.text:
                                    parts = para.text.split(":", 1)
                                    label_part = parts[0].strip()
                                    
                                    # Capture color before update
                                    original_color = None
                                    if para.runs:
                                        original_color = para.runs[0].font.color.rgb
                                    
                                    # Surgically update runs without clearing paragraph
                                    para.runs[0].text = f"{label_part} :"
                                    apply_run_style(para.runs[0], font_size=12, bold=True)
                                    if original_color:
                                        para.runs[0].font.color.rgb = original_color
                                    
                                    if len(para.runs) > 1:
                                        para.runs[1].text = f" {val}"
                                        apply_run_style(para.runs[1], font_size=12, bold=False)
                                        # Clear other runs if any
                                        for r_idx in range(2, len(para.runs)):
                                            para.runs[r_idx].text = ""
                                    else:
                                        v_run = para.add_run(f" {val}")
                                        apply_run_style(v_run, font_size=12, bold=False)
                        elif i + 1 < len(row.cells):
                            # Fill the target cell next to the label
                            target_cell = row.cells[i+1]
                            if target_cell.paragraphs:
                                first_p = target_cell.paragraphs[0]
                                original_color = None
                                if first_p.runs:
                                    original_color = first_p.runs[0].font.color.rgb
                                
                                safe_update_runs(first_p, str(val))
                                # Re-apply style to all runs
                                for r in first_p.runs:
                                    apply_run_style(r, font_size=12, bold=False)
                                    if original_color:
                                        r.font.color.rgb = original_color
                                
                                # Clear other paragraphs in the cell
                                for p_idx in range(1, len(target_cell.paragraphs)):
                                    for r in target_cell.paragraphs[p_idx].runs:
                                        r.text = ""
                        break

    # 2. Global Pass - Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Global Spacing Rule: Apply 1.5 spacing inside tables
                    paragraph.paragraph_format.line_spacing = 1.5
                    
                    if not paragraph.text.strip():
                        for run in paragraph.runs:
                            apply_run_style(run, font_size=12, bold=False)
                        continue
                    
                    # Use the non-destructive helper
                    process_label_value_paragraph(paragraph)
                    
                    # Ensure alignment and vertical alignment remain untouched
                    # Vertical alignment is cell-level, and paragraph.alignment was preserved in helper

    # 3. Global Pass - Main Body (Non-table paragraphs)
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
            
        text = paragraph.text
        # Inline Heading Detection (Non-table)
        is_heading = False
        if len(text) < 80:
            # Rules: ends with ':', OR is UPPERCASE, OR is Title Case, OR starts with list marker
            if (text.endswith(":") or 
                text.isupper() or 
                text.istitle() or 
                (len(text) > 2 and text[0].isalpha() and text[1] == ")")):
                is_heading = True
        
        apply_academic_formatting(paragraph, is_heading=is_heading)

    # 4. FINAL VALIDATION (MANDATORY)
    # Re-verify and Hard Lock "Journal Entry Topic"
    expected_topic = data.get("journal_topic", "").strip()
    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if "journal entry topic" in cell.text.lower():
                    # Check if it's label:value
                    if ":" in cell.text:
                        for p in cell.paragraphs:
                            if ":" in p.text:
                                parts = p.text.split(":", 1)
                                label = parts[0].strip()
                                current_val = parts[1].strip()
                                if current_val != expected_topic:
                                    # Force overwrite
                                    safe_update_runs(p, f"{label} : {expected_topic}")
                                    # Re-apply styling
                                    if p.runs:
                                        apply_run_style(p.runs[0], font_size=12, bold=True)
                                        if len(p.runs) > 1:
                                            # If safe_update_runs put everything in runs[0], 
                                            # we might need to split again if we want bold label
                                            # But safe_update_runs just updates runs[0].
                                            # Let's fix process_label_value_paragraph logic here too if needed.
                                            pass
                    elif i + 1 < len(row.cells):
                        target_cell = row.cells[i+1]
                        if target_cell.text.strip() != expected_topic:
                            if target_cell.paragraphs:
                                safe_update_runs(target_cell.paragraphs[0], expected_topic)
                                for r in target_cell.paragraphs[0].runs:
                                    apply_run_style(r, font_size=12, bold=False)

    doc.save(output_path)
    return output_path


def fill_free_writing(template_path: str, output_path: str, data: dict) -> str:
    # Load standard template directly
    doc = Document(template_path)

    # Standardize Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    student_details = data.get("student_details", {})
    
    mappings = {
        "name of the student": student_details.get("student_name", ""),
        "academic year": student_details.get("academic_year", ""),
        "student registration number": student_details.get("registration_number", ""),
        "registration number": student_details.get("registration_number", ""),
        "year & term": student_details.get("year_term", ""),
        "year and term": student_details.get("year_term", ""),
        "study level": student_details.get("study_level", ""),
        "class & section": student_details.get("class_section", ""),
        "class and section": student_details.get("class_section", ""),
        "course": student_details.get("course_name", ""),
        "instructor": student_details.get("instructor", ""),
        "assessment": student_details.get("assessment", ""),
        "date": data.get("date", ""),
    }

    # 1. Fill Student Metadata in Table 0
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        for row in t0.rows:
            for cell in row.cells:
                cell_text_lower = cell.text.lower().strip()
                for key, val in mappings.items():
                    if key in cell_text_lower:
                        if ":" in cell.text:
                            for para in cell.paragraphs:
                                if ":" in para.text:
                                    parts = para.text.split(":", 1)
                                    label_part = parts[0].strip()
                                    original_color = None
                                    if para.runs:
                                        original_color = para.runs[0].font.color.rgb
                                    
                                    para.runs[0].text = f"{label_part} :"
                                    apply_run_style(para.runs[0], font_size=12, bold=True)
                                    if original_color:
                                        para.runs[0].font.color.rgb = original_color
                                    
                                    if len(para.runs) > 1:
                                        para.runs[1].text = f" {val}"
                                        apply_run_style(para.runs[1], font_size=12, bold=False)
                                        for r_idx in range(2, len(para.runs)):
                                            para.runs[r_idx].text = ""
                                    else:
                                        v_run = para.add_run(f" {val}")
                                        apply_run_style(v_run, font_size=12, bold=False)
                        break

    # 2. Delete all content tables after the student metadata section (Table 0)
    tables_to_remove = list(doc.tables)[1:]
    for t in tables_to_remove:
        parent = t._element.getparent()
        if parent is not None:
            parent.remove(t._element)

    # 3. Clear Placeholder Body Paragraphs
    for p in list(doc.paragraphs):
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)

    # 4. Extract Dynamic Title and Sections List
    generated_data = data.get("generated_content", {})
    custom_title = ""
    sections_list = []

    if isinstance(generated_data, dict):
        custom_title = generated_data.get("title", "")
        sections_list = generated_data.get("sections", [])
    elif isinstance(generated_data, list):
        sections_list = generated_data

    title_text = custom_title or data.get("free_writing_topic", "")

    # Append Custom Title / Topic centered and bold
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(18)
    title_p.paragraph_format.line_spacing = 1.5
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(f"Topic: {title_text}")
    apply_run_style(title_run, font_size=16, bold=True)

    # 5. Append Generated Content with Dynamic Headings
    if isinstance(sections_list, list) and len(sections_list) > 0:
        for sec in sections_list:
            heading = sec.get("heading", "")
            content = sec.get("content", "")
            if not heading or not content:
                continue
            
            # Heading Paragraph
            h_p = doc.add_paragraph()
            h_p.paragraph_format.space_before = Pt(18)
            h_p.paragraph_format.space_after = Pt(6)
            h_p.paragraph_format.line_spacing = 1.5
            h_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            h_run = h_p.add_run(heading)
            apply_run_style(h_run, font_size=14, bold=True)
            
            # Content Paragraphs
            paragraphs = content.split("\n\n")
            for p_text in paragraphs:
                p_text = p_text.strip()
                if not p_text:
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.5
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                p_run = p.add_run(p_text)
                apply_run_style(p_run, font_size=12, bold=False)

    elif isinstance(generated_data, dict):
        # Fallback to key-value parsing if sections key is absent or sections_list is empty
        for key, val in generated_data.items():
            if key == "title" or key == "sections":
                continue
            nice_title = key.replace("_", " ").title()
            
            # Heading Paragraph
            h_p = doc.add_paragraph()
            h_p.paragraph_format.space_before = Pt(18)
            h_p.paragraph_format.space_after = Pt(6)
            h_p.paragraph_format.line_spacing = 1.5
            h_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            h_run = h_p.add_run(nice_title)
            apply_run_style(h_run, font_size=14, bold=True)
            
            # Content Paragraphs
            paragraphs = val.split("\n\n")
            for p_text in paragraphs:
                p_text = p_text.strip()
                if not p_text:
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.5
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                p_run = p.add_run(p_text)
                apply_run_style(p_run, font_size=12, bold=False)

    elif isinstance(generated_data, str):
        # Robust fallback for plain string content
        paragraphs = generated_data.split("\n\n")
        for p_text in paragraphs:
            p_text = p_text.strip()
            if not p_text:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            p_run = p.add_run(p_text)
            apply_run_style(p_run, font_size=12, bold=False)

    doc.save(output_path)
    return output_path


def fill_literature_survey(template_path: str, output_path: str, data: dict) -> str:
    # Load standard template directly
    doc = Document(template_path)

    # Standardize Margins slightly tighter to ensure single page cover safety
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    student_details = data.get("student_details", {})
    topic = data.get("topic", "")
    generated_data = data.get("generated_content", {})

    # Delete all tables in the document (since cover page and sections are fresh paragraphs)
    for table in list(doc.tables):
        parent = table._element.getparent()
        if parent is not None:
            parent.remove(table._element)

    # Clear Placeholder Body Paragraphs
    for p in list(doc.paragraphs):
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)

    # Helpers
    def add_run_style_local(run, font_size=14, bold=False, color=None):
        run.font.name = 'Times New Roman'
        run.font.size = Pt(font_size)
        run.bold = bold
        if color:
            run.font.color.rgb = color
        r = run._element.get_or_add_rPr()
        rFonts = r.get_or_add_rFonts()
        for attr in ['w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs']:
            rFonts.set(qn(attr), 'Times New Roman')

    def add_centered_paragraph(text, font_size=14, bold=False, italic=False, color=None, space_before=0, space_after=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        add_run_style_local(run, font_size=font_size, bold=bold, color=color)
        run.font.italic = italic
        return p

    def add_section_heading(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        add_run_style_local(run, font_size=14, bold=True)
        return p

    def add_article_heading(idx, title):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(f"{idx}. {title}")
        add_run_style_local(run, font_size=14, bold=True)
        return p

    def add_labeled_paragraph(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        l_run = p.add_run(f"{label} ")
        add_run_style_local(l_run, font_size=14, bold=True)
        
        v_run = p.add_run(str(value))
        add_run_style_local(v_run, font_size=14, bold=False)
        return p

    def add_body_paragraphs(content):
        if not content:
            return
        paragraphs = content.split("\n\n")
        for p_text in paragraphs:
            p_text = p_text.strip()
            if not p_text:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(p_text)
            add_run_style_local(run, font_size=14, bold=False)

    # 1. RENDER COVER PAGE
    add_centered_paragraph("Literature Survey", font_size=16, space_before=6, space_after=3)
    add_centered_paragraph("on", font_size=16, space_after=6)
    
    # Project Title (Bold, Red, 14pt)
    add_centered_paragraph(topic, font_size=14, bold=True, color=RGBColor(255, 0, 0), space_after=12)
    
    add_centered_paragraph("Submitted in partial fulfilment of the award of the", font_size=16, space_before=6, space_after=3)
    add_centered_paragraph("Bachelor of Technology", font_size=16, bold=True, space_after=3)
    add_centered_paragraph("in", font_size=16, space_after=3)
    
    dept_name = student_details.get("course_name", "[DEPARTMENT NAME]")
    if not dept_name.lower().startswith("department"):
        dept_name = f"Department of {dept_name}"
    add_centered_paragraph(dept_name, font_size=16, bold=True, space_after=12)
    
    add_centered_paragraph("By", font_size=16, space_after=6)
    
    student_names_str = student_details.get("student_name", "")
    student_rolls_str = student_details.get("registration_number", "")
    
    names = [n.strip() for n in student_names_str.split(",") if n.strip()]
    rolls = [r.strip() for r in student_rolls_str.split(",") if r.strip()]
    
    for idx in range(max(len(names), 1)):
        name = names[idx] if idx < len(names) else f"Student {idx+1}"
        roll = f"({rolls[idx]})" if idx < len(rolls) else ""
        add_centered_paragraph(f"{name} {roll}".strip(), font_size=16, space_after=2)
        
    add_centered_paragraph("Under the esteemed guidance of", font_size=16, space_before=12, space_after=6)
    
    guide_name = student_details.get("instructor", "[GUIDE NAME]")
    add_centered_paragraph(guide_name, font_size=16, bold=True, space_after=2)
    
    guide_class_sec = student_details.get("class_section", "")
    if "," in guide_class_sec:
        guide_designation, guide_department = [p.strip() for p in guide_class_sec.split(",", 1)]
    else:
        guide_designation = guide_class_sec
        guide_department = ""
        
    add_centered_paragraph(guide_designation, font_size=14, bold=True, space_after=2)
    if guide_department:
        add_centered_paragraph(guide_department, font_size=16, space_after=8)
    
    # Insert Extracted Aurora Logo
    import os
    logo_path = os.path.join(os.path.dirname(__file__), "..", "templates", "aurora_logo.jpeg")
    if os.path.exists(logo_path):
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.space_before = Pt(6)
        logo_p.paragraph_format.space_after = Pt(6)
        logo_p.paragraph_format.line_spacing = 1.15
        run = logo_p.add_run()
        run.add_picture(logo_path, width=Inches(1.2))
        
    univ_name = student_details.get("year_term", "[UNIVERSITY / COLLEGE NAME]")
    univ_loc = student_details.get("study_level", "[UNIVERSITY LOCATION]")
    
    add_centered_paragraph(univ_name.upper(), font_size=12, bold=True, space_after=2)
    add_centered_paragraph("(Deemed to be University)", font_size=12, bold=True, space_after=2)
    add_centered_paragraph(univ_loc, font_size=14, space_after=2)
    
    acad_year = student_details.get("academic_year", "[YEAR]")
    add_centered_paragraph(f"({acad_year})", font_size=14, space_before=3)
    
    # Page Break after Cover Page
    doc.add_page_break()

    # 2. RENDER INTRODUCTION
    add_section_heading("Introduction")
    add_body_paragraphs(generated_data.get("introduction", ""))

    # 3. RENDER OBJECTIVES
    add_section_heading("Objectives of the Project")
    obj_content = generated_data.get("objectives", "")
    obj_lines = [line.strip() for line in re.split(r'\r?\n', obj_content) if line.strip()]
    for line in obj_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(line)
        add_run_style_local(run, font_size=14, bold=False)

    # 4. RENDER LITERATURE REVIEW
    add_section_heading("Review of Articles")
    
    # Loop and render each article
    articles = generated_data.get("papers", [])
    for idx, art in enumerate(articles, 1):
        add_article_heading(idx, art.get("title", ""))
        add_labeled_paragraph("Title:", art.get("title", ""))
        add_labeled_paragraph("Authors:", art.get("authors", ""))
        add_labeled_paragraph("Year:", art.get("year", ""))
        add_labeled_paragraph("Abstract Summary:", art.get("abstract_summary", art.get("summary", "")))
        add_labeled_paragraph("Methodology:", art.get("methodology", ""))
        add_labeled_paragraph("Advantages:", art.get("advantages", ""))
        add_labeled_paragraph("Limitations:", art.get("limitations", ""))

    # 5. RENDER FINAL CONCLUSION
    add_section_heading("Conclusion")
    add_body_paragraphs(generated_data.get("conclusion", ""))

    # 6. RENDER REFERENCES
    if articles:
        add_section_heading("References")
        for idx, art in enumerate(articles, 1):
            ref_p = doc.add_paragraph()
            ref_p.paragraph_format.left_indent = Inches(0.5)
            ref_p.paragraph_format.first_line_indent = Inches(-0.5)
            ref_p.paragraph_format.space_after = Pt(6)
            ref_p.paragraph_format.line_spacing = 1.5
            ref_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            authors = art.get("authors", "")
            year = art.get("year", "")
            title = art.get("title", "")
            source = art.get("source", "")
            url = art.get("url", "")
            
            ref_text = f"[{idx}] {authors} ({year}). {title}."
            if source:
                ref_text += f" Retrieved via {source}."
            if url:
                ref_text += f" URL: {url}"
                
            run = ref_p.add_run(ref_text)
            add_run_style_local(run, font_size=12, bold=False)

    doc.save(output_path)
    return output_path

