import os
from utils.document_service import fill_free_writing
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def test_free_writing_with_template(template_name):
    template_path = f"templates/{template_name}"
    output_path = f"output/FreeWriting_Test_{template_name}"
    
    data = {
        "free_writing_topic": "The Ethics and Efficacy of Large Language Models in Academic Writing Assessments",
        "date": "2026-05-22",
        "student_details": {
            "student_name": "Test Student",
            "academic_year": "2025 - 2026",
            "registration_number": "REG123456",
            "year_term": "2nd Year, III Sem",
            "study_level": "UG",
            "class_section": "A",
            "course_name": "Computer Science",
            "instructor": "Dr. AI",
            "assessment": "Free Writing",
        },
        "generated_content": {
            "title": "Artificial Intelligence in Modern Science",
            "sections": [
                {
                    "heading": "Introduction",
                    "content": "This is the introduction paragraph. It should flow normally as an academic document without being inside any tables, boxes, or cells."
                },
                {
                    "heading": "Core Analysis of LLMs",
                    "content": "First paragraph of the core analysis. It needs to be justified, using Bookman Old Style 12pt, with 1.5 line spacing.\n\nSecond paragraph of the core analysis. Spacing between paragraphs should be naturally managed with space_after and space_before."
                },
                {
                    "heading": "Conclusion",
                    "content": "In conclusion, this test checks that free writing works seamlessly."
                }
            ]
        }
    }
    
    if not os.path.exists(template_path):
        print(f"Skipping test: {template_path} not found.")
        return

    final_path = fill_free_writing(template_path, output_path, data)
    print(f"Generated free writing doc for {template_name}: {final_path}")
    
    # Assertions on output
    doc = Document(final_path)
    
    # 1. Only Table 0 exists
    assert len(doc.tables) == 1, f"Expected exactly 1 table (Table 0), but found {len(doc.tables)}"
    
    # 2. Table 0 is filled correctly
    t0 = doc.tables[0]
    cell_texts = []
    for r in t0.rows:
        for c in r.cells:
            cell_texts.append(c.text.lower())
    
    assert any("test student" in t for t in cell_texts), "Student name not found in Table 0"
    assert any("reg123456" in t for t in cell_texts), "Registration number not found in Table 0"
    
    # 3. Headings and Paragraphs check
    # Check that we have a centered title: Topic: Artificial Intelligence in Modern Science
    title_p = None
    for p in doc.paragraphs:
        if "topic:" in p.text.lower():
            title_p = p
            break
            
    assert title_p is not None, "Title paragraph 'Topic: ...' not found in body"
    assert title_p.alignment == WD_ALIGN_PARAGRAPH.CENTER, "Title paragraph is not centered"
    
    # Check that runs have Bookman Old Style
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        for run in p.runs:
            assert run.font.name == 'Bookman Old Style', f"Font is not Bookman Old Style on run: {run.text}"
            
    print(f"Test for {template_name} passed successfully!")

if __name__ == "__main__":
    os.makedirs("output", exist_ok=True)
    test_free_writing_with_template("standard_assignment.docx")
    test_free_writing_with_template("free_writing_template.docx")
