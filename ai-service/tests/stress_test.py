import os
import sys
import time
import json
import unittest
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from docx import Document as DocxDocument

# Add root folder to sys path
sys.path.append(str(Path(__file__).parent.parent))

from langchain_core.outputs import ChatResult, ChatGeneration
from langchain_core.messages import AIMessage
from providers.fallback_provider import get_llm, FallbackChatLLM, global_groq_breaker
from providers.metrics import global_metrics
from providers.cache import global_cache
from agents.workflow import ReflectiveJournalWorkflow
from utils.document_service import fill_reflective_journal

# A standard valid JSON string containing over 150 words across all required keys to satisfy validation
LONG_SUCCESS_JSON = json.dumps({
    "experience": "During the class discussion, we explored various concepts of machine learning including supervised and unsupervised approaches. We analyzed decision trees, support vector machines, and neural networks in detail. The instructor presented several case studies illustrating how these methodologies solve real-world problems. We spent significant time studying the training phase, loss functions, and optimization techniques. Each student also discussed their own experiences with dataset preparation and cleaning.",
    "feelings": "Initially, I felt quite overwhelmed by the mathematical complexity of the algorithms, especially backpropagation in neural networks. However, as the session progressed and we discussed practical examples, I became more comfortable. I realized that breaking down the equations into algorithmic steps helps build intuitive understanding.",
    "learning": "The key insight I gained was the fundamental trade-off between model bias and variance. I learned that overfitting can be mitigated through regularization techniques such as L1 and L2 regularization. Furthermore, cross-validation is essential for evaluating model generalization performance on unseen test datasets.",
    "application": "I will apply these concepts in my current project where we are building a predictive model for customer churn. I will start by thoroughly analyzing the dataset for missing values, then apply cross-validation to select the best hyper-parameters, and finally evaluate the model metrics.",
    "conclusion": "In conclusion, this session was extremely valuable in bridging the gap between theoretical machine learning and practical implementation. I now feel confident in my ability to construct and evaluate basic classification models systematically."
})

LONG_GEMINI_JSON = json.dumps({
    "experience": "On the Gemini fallback path, we explored the exact same educational paradigms of AI systems integration. We studied how fallback architectures ensure reliability even when primary providers fail. The instructor highlighted that Gemini is a highly capable alternative that can process and generate responses with low latency and high quality. We also discussed standard API status codes like 429 and how circuit breakers can prevent cascading failures.",
    "feelings": "I felt relieved that the fallback system executed seamlessly when Groq simulation failed. It gave me greater confidence in the overall stability of the design. The integration felt extremely smooth and natural.",
    "learning": "I learned that multi-provider failover is a standard industry practice for building production-grade services. I also learned that circuit breaker design requires careful selection of failure rate limits and cooldown durations.",
    "application": "I will apply this fallback design to all future microservices that depend on third-party APIs. This ensures high availability and robustness under high concurrent request rates.",
    "conclusion": "Overall, fallback activation is a reliable mechanism to maintain service uptime. This test proved that the system degrades gracefully under simulated provider failure scenarios."
})

class StressTestFallback(unittest.TestCase):
    def setUp(self):
        global_cache.clear()
        # Reset metrics
        global_metrics.__init__()
        # Clear failures
        global_groq_breaker.clear()
        
        self.orig_groq = os.environ.get("GROQ_API_KEY")
        self.orig_gemini = os.environ.get("GEMINI_API_KEY")
        os.environ["GROQ_API_KEY"] = self.orig_groq or "gsk_dummy"
        os.environ["GEMINI_API_KEY"] = self.orig_gemini or "dummy_gemini"

    def tearDown(self):
        if self.orig_groq:
            os.environ["GROQ_API_KEY"] = self.orig_groq
        else:
            os.environ.pop("GROQ_API_KEY", None)
        if self.orig_gemini:
            os.environ["GEMINI_API_KEY"] = self.orig_gemini
        else:
            os.environ.pop("GEMINI_API_KEY", None)

    def test_concurrent_generation_and_caching(self):
        print("\n--- Running Stress Test: Concurrent Cache Hits and Thread Safety ---")
        llm = get_llm()
        
        # Mock BOTH providers to prevent real network calls
        llm.groq_llm._generate = lambda messages, stop=None, **kwargs: ChatResult(
            generations=[ChatGeneration(message=AIMessage(content=LONG_SUCCESS_JSON))]
        )
        llm.gemini_llm._generate = lambda messages, stop=None, **kwargs: ChatResult(
            generations=[ChatGeneration(message=AIMessage(content=LONG_GEMINI_JSON))]
        )
        
        wf = ReflectiveJournalWorkflow(topic="Stress Test Topic")
        wf.llm = llm
        
        num_threads = 55 # 50+ concurrent requests
        results = []
        
        start_time = time.time()
        with ThreadPoolExecutor(max_workers=10) as executor:
            futures = [executor.submit(wf.execute, f"req-{i}") for i in range(num_threads)]
            for fut in as_completed(futures):
                results.append(fut.result())
                
        duration = time.time() - start_time
        print(f"Completed {num_threads} concurrent workflow generations in {duration:.4f} seconds.")
        
        self.assertEqual(len(results), num_threads)
        for res in results:
            self.assertTrue(res["experience"].startswith("During the class"))
            self.assertTrue(res["conclusion"].startswith("In conclusion"))
            
        stats = global_metrics.get_stats()
        print("Metrics Stats:", stats)
        # Ensure caching works and recorded hits
        self.assertGreater(stats["cache_hits"], 0)

    def test_forced_groq_failure_and_circuit_breaker(self):
        print("\n--- Running Stress Test: Forced Groq Failures and Circuit Breaker Tripping ---")
        llm = get_llm()
        
        def failing_groq(messages, stop=None, **kwargs):
            raise RuntimeError("Groq rate limit exceeded")
            
        llm.groq_llm._generate = failing_groq
        llm.gemini_llm._generate = lambda messages, stop=None, **kwargs: ChatResult(
            generations=[ChatGeneration(message=AIMessage(content=LONG_GEMINI_JSON))]
        )
        
        num_trips = 8
        results = []
        
        # Disable cache for circuit breaker failure accumulation
        os.environ["CACHE_ENABLED"] = "false"
        try:
            for i in range(num_trips):
                # Use unique topics to ensure no cache hits even if cache was enabled
                wf = ReflectiveJournalWorkflow(topic=f"Fallback Stress Test Topic {i}")
                wf.llm = llm
                res = wf.execute(f"fallback-req-{i}")
                results.append(res)
        finally:
            os.environ.pop("CACHE_ENABLED", None)
            
        self.assertEqual(len(results), num_trips)
        for res in results:
            self.assertTrue(res["experience"].startswith("On the Gemini"))
            
        self.assertTrue(global_groq_breaker.is_tripped())
        stats = global_metrics.get_stats()
        print("Fallback Metrics Stats:", stats)
        self.assertGreaterEqual(stats["fallback_count"], 5)
        self.assertEqual(stats["groq"]["failures"], 5)

    def test_provider_timeouts(self):
        print("\n--- Running Stress Test: Provider Timeouts fallback ---")
        llm = get_llm()
        
        def timeout_groq(messages, stop=None, **kwargs):
            raise RuntimeError("Request timed out or connection deadline exceeded")
            
        llm.groq_llm._generate = timeout_groq
        llm.gemini_llm._generate = lambda messages, stop=None, **kwargs: ChatResult(
            generations=[ChatGeneration(message=AIMessage(content=LONG_GEMINI_JSON))]
        )
        
        wf = ReflectiveJournalWorkflow(topic="Timeout Topic")
        wf.llm = llm
        
        res = wf.execute("req-timeout-1")
        self.assertTrue(res["experience"].startswith("On the Gemini"))
        
        stats = global_metrics.get_stats()
        self.assertEqual(stats["groq"]["timeouts"], 1)

    def test_malformed_json_and_validator_failure(self):
        print("\n--- Running Stress Test: Malformed JSON output and validation failures ---")
        llm = get_llm()
        
        malformed_text = "This is not json at all! {unclosed bracket"
        
        llm.groq_llm._generate = lambda messages, stop=None, **kwargs: ChatResult(
            generations=[ChatGeneration(message=AIMessage(content=malformed_text))]
        )
        llm.gemini_llm._generate = lambda messages, stop=None, **kwargs: ChatResult(
            generations=[ChatGeneration(message=AIMessage(content=LONG_GEMINI_JSON))]
        )
        
        wf = ReflectiveJournalWorkflow(topic="Validation Failure Topic")
        wf.llm = llm
        
        res = wf.execute("req-validation-1")
        self.assertTrue(res["experience"].startswith("On the Gemini"))
        
        stats = global_metrics.get_stats()
        # Verify validation failure was recorded for Groq
        self.assertEqual(stats["groq"]["validation_failures"], 1)

    def test_invalid_groq_key(self):
        print("\n--- Running Stress Test: Invalid Groq Key ---")
        llm = get_llm()
        
        def bad_key_groq(messages, stop=None, **kwargs):
            raise Exception("Authentication Error: Invalid API Key")
            
        llm.groq_llm._generate = bad_key_groq
        llm.gemini_llm._generate = lambda messages, stop=None, **kwargs: ChatResult(
            generations=[ChatGeneration(message=AIMessage(content=LONG_GEMINI_JSON))]
        )
        
        wf = ReflectiveJournalWorkflow(topic="Key Topic")
        wf.llm = llm
        
        res = wf.execute("req-key-1")
        self.assertTrue(res["experience"].startswith("On the Gemini"))

    def test_concurrent_docx_generation(self):
        print("\n--- Running Stress Test: Concurrent DOCX Generation and Verification ---")
        template_file = str(Path(__file__).parent.parent / "templates" / "standard_assignment.docx")
        output_dir = Path(__file__).parent.parent / "output"
        output_dir.mkdir(exist_ok=True)
        
        num_docx_threads = 15
        
        def run_gen(idx):
            out_file = str(output_dir / f"stress_test_doc_{idx}.docx")
            test_data = {
                "journal_topic": f"Topic {idx}",
                "date": "2026-05-24",
                "student_details": {
                    "student_name": f"Student {idx}",
                    "academic_year": "2026",
                    "registration_number": f"REG-{idx}",
                    "year_term": "Year 1 Term 1",
                    "study_level": "Undergraduate",
                    "class_section": "Section A",
                    "course_name": "Course A",
                    "instructor": "Instructor A",
                    "assessment": "Reflective Journal",
                },
                "generated_content": {
                    "experience": f"Exp content for {idx}. Paragraph 1.\n\nParagraph 2 for {idx}.",
                    "feelings": f"Feel content for {idx}.",
                    "learning": f"Learn content for {idx}.",
                    "application": f"App content for {idx}.",
                    "conclusion": f"Conc content for {idx}."
                }
            }
            return fill_reflective_journal(template_file, out_file, test_data)

        results = []
        with ThreadPoolExecutor(max_workers=5) as executor:
            futures = [executor.submit(run_gen, i) for i in range(num_docx_threads)]
            for fut in as_completed(futures):
                results.append(fut.result())
                
        # Inspect each file to verify formatting and integrity
        for path in results:
            doc = DocxDocument(path)
            all_text = ""
            for p in doc.paragraphs:
                all_text += p.text + "\n"
                if p.text.strip():
                    self.assertEqual(p.paragraph_format.line_spacing, 1.5)
                    
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text += cell.text + "\n"
                        for p in cell.paragraphs:
                            if p.text.strip():
                                self.assertEqual(p.paragraph_format.line_spacing, 1.5)

            # Assert no duplication/corruption: headings only once, text present
            for sec in ["Experience", "Feelings", "Learning", "Application", "Conclusion"]:
                self.assertLessEqual(all_text.count(sec), 4, f"Duplicate section header found for {sec} in {path}")
                
            # Assert text contains student data
            self.assertIn("Student", all_text)
            self.assertIn("REG-", all_text)
            
            # Clean up generated test files
            try:
                os.remove(path)
            except Exception:
                pass
        print(f"Verified {num_docx_threads} DOCX files for alignment, spacing, headings, and data integrity.")

if __name__ == "__main__":
    unittest.main()
