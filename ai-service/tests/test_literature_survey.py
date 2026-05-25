import unittest
from unittest.mock import patch, MagicMock
from pathlib import Path
from docx import Document

from services.literature.models import Paper
from services.literature.normalizer import validate_and_normalize_paper, deduplicate_papers
from services.literature.semantic_scholar import SemanticScholarProvider
from services.literature.arxiv_provider import ArxivProvider
from services.literature.search_manager import SearchManager
from agents.literature_survey_workflow import LiteratureSurveyWorkflow
from utils.document_service import fill_literature_survey


class TestLiteratureSurveyPipeline(unittest.TestCase):
    def setUp(self):
        # Patch get_llm to avoid EnvironmentError when run without keys
        self.get_llm_patcher = patch("agents.workflow.get_llm")
        self.mock_get_llm = self.get_llm_patcher.start()
        self.mock_llm_instance = MagicMock()
        self.mock_get_llm.return_value = self.mock_llm_instance

        self.base_dir = Path(__file__).parent.parent
        self.tmp_dir = self.base_dir / "output"
        self.tmp_dir.mkdir(parents=True, exist_ok=True)
        self.template_path = self.tmp_dir / "survey_template_test.docx"
        self.output_path = self.tmp_dir / "survey_output_test.docx"

        # Generate a test document with a metadata table at index 0
        doc = Document()
        table = doc.add_table(rows=6, cols=2)
        table.cell(0, 0).text = "Name of the Student"
        table.cell(1, 0).text = "Student Registration Number"
        table.cell(2, 0).text = "Academic Year"
        table.cell(3, 0).text = "Course"
        table.cell(4, 0).text = "Instructor"
        table.cell(5, 0).text = "Date of Submission"
        
        doc.add_paragraph("Literature Survey Topic placeholder")
        doc.save(str(self.template_path))

        # Sample paper payloads (direct Direct direct keys format returned by clean Semantic Scholar)
        self.raw_ss_paper = {
            "paperId": "ss_123",
            "title": "Deep Learning for Grading University Essays",
            "abstract": "This paper presents a novel neural network architecture for grading essays. The model utilizes BERT embeddings and achieves high correlation with human assessors.",
            "authors": [{"name": "Atherton John"}, {"name": "Smith Robert"}],
            "year": 2024,
            "venue": "IEEE Transactions on Education",
            "citationCount": 15,
            "externalIds": {"DOI": "10.1109/TE.2024.12345"},
            "url": "https://api.semanticscholar.org/ss_123"
        }

        # Similar title for testing deduplication similarity matching
        self.raw_arxiv_paper = {
            "id": "arxiv_456",
            "title": "Deep Learning for Grading University Essays!",
            "abstract": "We evaluate an alternative neural network design for scoring essays. Our findings show that simple LSTM models perform on par with transformers for short texts.",
            "authors": [{"name": "John Atherton"}, {"name": "R. Smith"}],
            "year": 2024,
            "venue": "ArXiv",
            "citationCount": 2,
            "doi": None,
            "url": "http://arxiv.org/abs/2405.0001"
        }

    def tearDown(self):
        self.get_llm_patcher.stop()

    def test_paper_normalization_and_validation(self):
        # Test valid normalization
        paper = validate_and_normalize_paper(self.raw_ss_paper, "semantic_scholar")
        self.assertIsNotNone(paper)
        self.assertEqual(paper.title, "Deep Learning for Grading University Essays")
        self.assertEqual(paper.authors, ["Atherton John", "Smith Robert"])
        self.assertEqual(paper.year, 2024)
        self.assertEqual(paper.doi, "10.1109/TE.2024.12345")
        self.assertEqual(paper.source, "semantic_scholar")

        # Test malformed paper rejection (e.g. missing title)
        malformed = self.raw_ss_paper.copy()
        malformed["title"] = ""
        paper_rejected = validate_and_normalize_paper(malformed, "semantic_scholar")
        self.assertIsNone(paper_rejected)

    def test_deduplication_engine(self):
        p1 = validate_and_normalize_paper(self.raw_ss_paper, "semantic_scholar")
        p2 = validate_and_normalize_paper(self.raw_arxiv_paper, "arxiv")
        
        papers_list = [p1, p2]
        
        # Test title Jaccard similarity deduplication (similarity > 0.85 should merge these titles)
        deduped = deduplicate_papers(papers_list)
        self.assertEqual(len(deduped), 1)
        self.assertEqual(deduped[0].id, p1.id)

    @patch("services.literature.semantic_scholar.httpx.Client")
    def test_semantic_scholar_provider(self, mock_client_cls):
        mock_client = MagicMock()
        mock_response = MagicMock()
        mock_response.status_code = 200
        mock_response.json.return_value = {"data": [self.raw_ss_paper]}
        mock_client.get.return_value = mock_response
        mock_client_cls.return_value.__enter__.return_value = mock_client

        provider = SemanticScholarProvider()
        results = provider.search("deep learning essays")
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]["id"], "ss_123")

    @patch("services.literature.arxiv_provider.httpx.Client")
    def test_arxiv_xml_defensive_parsing(self, mock_client_cls):
        # Mocking arXiv atom feed response XML content
        mock_xml = """<?xml version="1.0" encoding="utf-8"?>
        <feed xmlns="http://www.w3.org/2005/Atom">
          <entry>
            <id>http://arxiv.org/abs/2405.0001</id>
            <title>Deep Learning for Grading University Essays</title>
            <summary>We evaluate an alternative neural network design for scoring essays.</summary>
            <author><name>John Atherton</name></author>
            <author><name>R. Smith</name></author>
            <published>2024-05-10T12:00:00Z</published>
            <link href="http://arxiv.org/abs/2405.0001" rel="alternate" type="text/html"/>
          </entry>
        </feed>
        """
        mock_client = MagicMock()
        mock_response = MagicMock()
        mock_response.status_code = 200
        mock_response.text = mock_xml
        mock_client.get.return_value = mock_response
        mock_client_cls.return_value.__enter__.return_value = mock_client

        provider = ArxivProvider()
        results = provider.search("deep learning essays")
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]["title"], "Deep Learning for Grading University Essays")

    @patch.object(SemanticScholarProvider, "search")
    @patch.object(ArxivProvider, "search")
    def test_search_manager_fallback(self, mock_arxiv_search, mock_ss_search):
        # Semantic scholar returns 0 results, forcing ArXiv fallback
        mock_ss_search.return_value = []
        mock_arxiv_search.return_value = [self.raw_arxiv_paper]

        manager = SearchManager()
        results = manager.search("neural essays")
        
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0].source, "arxiv")
        mock_arxiv_search.assert_called_once()

    def test_workflow_citation_anti_hallucination_validation(self):
        p1 = validate_and_normalize_paper(self.raw_ss_paper, "semantic_scholar")
        p2 = validate_and_normalize_paper(self.raw_arxiv_paper, "arxiv")
        
        # p1 authors: Atherton John, Smith Robert
        # p2 authors: John Atherton, R. Smith
        # Surnames: john, robert, atherton, smith
        workflow = LiteratureSurveyWorkflow(
            topic="Essay grading neural networks",
            selected_papers=[p1.dict() if hasattr(p1, "dict") else p1.model_dump(), p2.dict() if hasattr(p2, "dict") else p2.model_dump()]
        )

        # 1. Valid citation structure
        valid_payload = {
            "introduction": "Atherton et al. (2024) introduced key grading methods.",
            "objectives": "Analyze modern grading techniques.",
            "advantages": "Robust scores.",
            "disadvantages": "High computational cost.",
            "paper_summaries": "Atherton and Smith (2024) designed a neural scorer. John et al. (2024) checked another approach.",
            "conclusion": "Grading is progressing."
        }
        is_valid, err = workflow._validate_citations(valid_payload)
        self.assertTrue(is_valid, f"Expected valid citation check: {err}")

        # 2. Fabricated citation structure (contains author "Miller (2024)" which is not in our list of papers)
        invalid_payload = {
            "introduction": "Atherton et al. (2024) introduced grading methods, but Miller (2024) contradicted this.",
            "objectives": "Compare grading methods.",
            "advantages": "Robust scores.",
            "disadvantages": "High computational cost.",
            "paper_summaries": "Atherton (2024) designed a neural scorer.",
            "conclusion": "Grading is progressing."
        }
        is_valid, err = workflow._validate_citations(invalid_payload)
        self.assertFalse(is_valid)
        self.assertIn("Citation validation failed", err)

    def test_fill_literature_survey_docx_rendering(self):
        p1 = validate_and_normalize_paper(self.raw_ss_paper, "semantic_scholar")
        p2 = validate_and_normalize_paper(self.raw_arxiv_paper, "arxiv")
        
        data = {
            "topic": "Neural essay scoring in higher education",
            "date": "2026-05-25",
            "student_details": {
                "student_name": "Vignesh Kumar",
                "academic_year": "2025 - 2026",
                "registration_number": "241U1R2089",
                "year_term": "2nd Year, III Sem",
                "study_level": "UG",
                "class_section": "A",
                "course_name": "Artificial Intelligence",
                "instructor": "Prof. Smith",
                "assessment": "Literature Survey Assessment",
            },
            "generated_content": {
                "introduction": "Introduction content text. Discusses historical scoring models.",
                "objectives": "Objectives content text. Examines specific neural performance.",
                "advantages": "Advantages content text. Synthesizes key benefits.",
                "disadvantages": "Disadvantages content text. Discusses cost models.",
                "paper_summaries": "Atherton (2024) designed BERT models. Smith (2024) evaluated LSTM variations.",
                "conclusion": "Conclusion content text. Outlines future research vectors.",
                "papers": [
                    {
                        "title": p1.title,
                        "authors": ", ".join(p1.authors),
                        "year": str(p1.year),
                        "summary": "BERT-based essay scoring model.",
                        "url": p1.url,
                        "source": p1.source
                    },
                    {
                        "title": p2.title,
                        "authors": ", ".join(p2.authors),
                        "year": str(p2.year),
                        "summary": "LSTM-based essay scoring model comparison.",
                        "url": p2.url,
                        "source": p2.source
                    }
                ]
            }
        }

        output = fill_literature_survey(
            template_path=str(self.template_path),
            output_path=str(self.output_path),
            data=data
        )

        self.assertTrue(Path(output).exists())
        
        # Verify rendered document contents
        doc = Document(output)
        para_text = "\n".join(p.text for p in doc.paragraphs)
        table_text = "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        all_text = f"{para_text}\n{table_text}"

        self.assertIn("Vignesh Kumar", all_text)
        self.assertIn("241U1R2089", all_text)
        self.assertIn("Artificial Intelligence", all_text)
        self.assertIn("Literature Survey", all_text)
        self.assertIn("Introduction", all_text)
        self.assertIn("Review of Articles", all_text)
        self.assertIn("References", all_text)
        self.assertIn("Atherton John", all_text)
        self.assertIn("https://api.semanticscholar.org/ss_123", all_text)


if __name__ == "__main__":
    unittest.main()
